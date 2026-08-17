import io

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from app.schemas import ExportRequest

router = APIRouter(tags=["export"])


def _build_docx(req: ExportRequest) -> io.BytesIO:
    doc = DocxDocument()
    doc.add_heading(req.question, level=1)
    doc.add_paragraph(req.answer)

    if req.sources:
        doc.add_heading("Sources", level=2)
        for s in req.sources:
            label = s.title + (f" — {s.author}" if s.author else "")
            doc.add_paragraph(label, style="List Bullet")
            doc.add_paragraph(s.excerpt, style="Intense Quote")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_pdf(req: ExportRequest) -> io.BytesIO:
    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(buf, pagesize=LETTER)
    story = [Paragraph(req.question, styles["Title"]), Spacer(1, 12)]
    for para in req.answer.split("\n"):
        if para.strip():
            story.append(Paragraph(para, styles["BodyText"]))
            story.append(Spacer(1, 8))

    if req.sources:
        story.append(Spacer(1, 16))
        story.append(Paragraph("Sources", styles["Heading2"]))
        for s in req.sources:
            label = s.title + (f" — {s.author}" if s.author else "")
            story.append(Paragraph(f"<b>{label}</b>", styles["BodyText"]))
            story.append(Paragraph(s.excerpt, styles["Italic"]))
            story.append(Spacer(1, 8))

    doc.build(story)
    buf.seek(0)
    return buf


@router.post("/export")
def export_answer(req: ExportRequest):
    if req.format == "docx":
        buf = _build_docx(req)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=answer.docx"},
        )
    elif req.format == "pdf":
        buf = _build_pdf(req)
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=answer.pdf"},
        )
    else:
        raise HTTPException(status_code=400, detail="format must be 'pdf' or 'docx'")
